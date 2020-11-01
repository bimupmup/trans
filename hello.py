
import random
import spacy
import pyttsx3
from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from googlesearch import search
import requests, sys, webbrowser, bs4
from chatterbot import ChatBot
from chatterbot.trainers import ChatterBotCorpusTrainer
import googletrans
from googletrans import Translator
import docx
import os
from tqdm import tqdm
#trainer.train("chatterbot.corpus.wiki")

def chat_bot(message):
    return str(english_bot.get_response(message))

def run_bot():
	nlp = spacy.load('en')
	english_bot=ChatBot("Chatterbot",storage_adapter="chatterbot.storage.SQLStorageAdapter")
	trainer = ChatterBotCorpusTrainer(english_bot)
	trainer.train("chatterbot.corpus.english")
	while True:
		message=input()
		bot_say=chat_bot(message)
		print(bot_say)

#voices = engine.getProperty('voices')
#for voice in voices:
#	print("Voice:")
#	print("- ID: %s" % voice.id)
#	print("- Name: %s" % voice.name)
#	print("- Languages: %s" % voice.languages)
#	print("- Gender %s" % voice.gender)
#	print("- Age %s" % voice.age)
#you="hello"
#if you=="":
#	robot_say = "I can't hear you"
#elif you=="hello":
#	robot_say = "hello mtt"
#else:
#	robot_say = "I'm fine"
#while True:
#	text = input("cc \n")  input way
#	print(text)

#urL="www.facebook.com"
#chrome_path="C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe"
#address=webbrowser.register('chrome',webbrowser.BackgroundBrowser(chrome_path),1)
#webbrowser.get("chrome").open_new(address)
#webbrowser.open_new(urL)



#response = googlesearch().search("something")
#for result in response.results:
#    print("Title: " + result.title)
#   print("Content: " + rescmdult.getText())
str="open animal"
print(str)
#driver=webdriver.Chrome(chrome_path)
#driver.get("http://www.google.com")
#webbrowser.get("C:/Drivers/storage/chromedriver.exe").open_new("google.com")
#webbrowser.register('chrome', None,webbrowser.BackgroundBrowser(chrome_path))
def open_googleMap():
    chrome_path = "C:/Drivers/storage/chromedriver.exe"
    urL1 = "file:///C:/Users/Dell/Desktop/localWeb1.htm"
    urL2 = "https:/www.google.com/maps"
    #webbrowser.open_new(urL)
    browser= webdriver.Chrome(chrome_path)
    browser.get(urL1)
	
    loginButton = browser.find_element_by_css_selector\
    ("button.cancelbtn")
    loginButton2 =  browser.find_element_by_xpath\
    ("//button[@type='submit'][@name='login']")
    name = browser.find_element_by_xpath("//div[@class='container']/input[1]")
    name.send_keys("mtt")
    passWord = browser.find_element_by_xpath("//div[@class='container']/input[2]")
    passWord.send_keys("cc")

def search_engine(key):
    arr = key.split()
    subAdress = ""
    for i in range(0,len(arr)):
        subAdress=subAdress+arr[i]
        if(i<len(arr)-1):
            subAdress=subAdress+"+"
    for i in range(1,random.randint(2,5)):
        webbrowser.open('https://google.com/search?q=' + subAdress)
    print(subAdress)

def google_trans(text):
	translator = Translator()
	type = translator.detect(text)
	print("Traslated from vi",translator.translate(text))

google_trans("Xin chào, tôi là mtt")

def translate_tool(link):
	doc = docx.Document(os.path.join(link))
	docWritten = docx.Document()
	i = 0
	translator = Translator()
	numRow = len(doc.paragraphs)
	for i in tqdm (range(0,numRow)):
		sentence = translator.translate(doc.paragraphs[i].text)
		docWritten.add_paragraph(sentence.text)
	docWritten.save("261-265E.docx")
translate_tool("C:\\Users\\Administrator\\Downloads\\261-265V.docx")
#while True:
#	print("cc")